from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "gdd"
ASSET_DIR = OUT_DIR / "assets"
GENERATED_DIR = ASSET_DIR / "generated"
OUTPUT = OUT_DIR / "Drowned_Game_Design_Document.docx"

NAVY = "071A2B"
DEEP = "0B2840"
TEAL = "0E6F78"
CYAN = "3DD6D0"
GOLD = "F2B84B"
CORAL = "EF6A5B"
PURPLE = "7656A5"
INK = "17212B"
MUTED = "5A6874"
PALE = "EAF5F3"
PAPER = "F7F4EB"
WHITE = "FFFFFF"
LIGHT = "D9E7E6"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                edge.set(qn(f"w:{key}"), str(edge_data[key]))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Aptos", size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_hyperlink(paragraph, text: str, url: str, color=TEAL, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def fit_font(size: int, bold=False):
    windows_fonts = Path("C:/Windows/Fonts")
    preferred = windows_fonts / ("arialbd.ttf" if bold else "arial.ttf")
    if preferred.exists():
        return ImageFont.truetype(str(preferred), size=size)
    for candidate in ("DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf", "Arial"):
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_diamond(draw, cx, cy, w, h, fill, outline, width=3):
    pts = [(cx, cy - h / 2), (cx + w / 2, cy), (cx, cy + h / 2), (cx - w / 2, cy)]
    draw.polygon(pts, fill=fill)
    draw.line(pts + [pts[0]], fill=outline, width=width, joint="curve")


def build_visual_assets() -> None:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    # Cover artwork: submerged geometry and a lone diver silhouette.
    img = Image.new("RGB", (1600, 900), f"#{NAVY}")
    d = ImageDraw.Draw(img)
    for y in range(900):
        t = y / 899
        c = tuple(int(a * (1 - t) + b * t) for a, b in zip((7, 35, 56), (5, 12, 26)))
        d.line((0, y, 1600, y), fill=c)
    for x, y, r, a in [(210, 130, 20, 0), (1280, 210, 34, 0), (970, 110, 13, 0), (410, 250, 9, 0)]:
        d.ellipse((x-r, y-r, x+r, y+r), outline=f"#{CYAN}", width=3)
    d.polygon([(0, 720), (280, 560), (520, 720), (760, 500), (1080, 720), (1360, 530), (1600, 680), (1600, 900), (0, 900)], fill="#0B2736")
    d.line([(80, 735), (330, 585), (520, 730)], fill="#2C7778", width=6)
    d.line([(650, 725), (830, 545), (1050, 730)], fill="#2C7778", width=6)
    d.rectangle((755, 420, 845, 590), fill="#102A3A", outline=f"#{GOLD}", width=4)
    d.ellipse((760, 345, 840, 435), fill="#102A3A", outline=f"#{GOLD}", width=4)
    d.line((775, 585, 725, 700), fill=f"#{GOLD}", width=15)
    d.line((825, 585, 875, 700), fill=f"#{GOLD}", width=15)
    d.line((758, 465, 675, 555), fill=f"#{GOLD}", width=13)
    d.line((842, 465, 930, 535), fill=f"#{GOLD}", width=13)
    d.arc((780, 300, 910, 430), 210, 360, fill=f"#{CYAN}", width=8)
    d.text((105, 82), "DROWNED", font=fit_font(112, True), fill=f"#{WHITE}")
    d.text((112, 205), "GAME DESIGN DOCUMENT", font=fit_font(34, True), fill=f"#{GOLD}")
    img.save(GENERATED_DIR / "cover.png", quality=95)

    # Core loop.
    img = Image.new("RGB", (1600, 640), f"#{PAPER}")
    d = ImageDraw.Draw(img)
    labels = [
        ("ESCOLHER", "personagem e estilo"),
        ("EXPLORAR", "ruinas e rotas"),
        ("COMBATER", "posicao e habilidade"),
        ("DESCOBRIR", "dialogos e escolhas"),
        ("PROGREDIR", "novas areas e builds"),
    ]
    cols = [CORAL, TEAL, PURPLE, GOLD, CYAN]
    centers = [(800 + 490 * math.cos(-math.pi/2 + i*2*math.pi/5), 320 + 220 * math.sin(-math.pi/2 + i*2*math.pi/5)) for i in range(5)]
    for i, ((title, sub), col) in enumerate(zip(labels, cols)):
        x, y = centers[i]
        d.ellipse((x-150, y-64, x+150, y+64), fill=f"#{col}", outline=f"#{NAVY}", width=4)
        bbox = d.textbbox((0, 0), title, font=fit_font(27, True))
        d.text((x-(bbox[2]-bbox[0])/2, y-36), title, font=fit_font(27, True), fill=f"#{NAVY}")
        bbox = d.textbbox((0, 0), sub, font=fit_font(18))
        d.text((x-(bbox[2]-bbox[0])/2, y+8), sub, font=fit_font(18), fill=f"#{INK}")
        nx, ny = centers[(i+1) % 5]
        ax = x + (nx-x)*0.34
        ay = y + (ny-y)*0.34
        bx = x + (nx-x)*0.66
        by = y + (ny-y)*0.66
        d.line((ax, ay, bx, by), fill=f"#{NAVY}", width=7)
    d.ellipse((655, 245, 945, 395), fill=f"#{NAVY}")
    d.text((713, 278), "CICLO", font=fit_font(35, True), fill=f"#{WHITE}")
    d.text((690, 325), "de uma expedicao", font=fit_font(23), fill=f"#{CYAN}")
    img.save(GENERATED_DIR / "core_loop.png", quality=95)

    # Character placeholder cards.
    img = Image.new("RGB", (1600, 720), f"#{PAPER}")
    d = ImageDraw.Draw(img)
    cards = [
        ("QUEBRA-MAR", "Vanguarda", CORAL, "GOLPE", "curto alcance"),
        ("VIGIA", "Atirador", GOLD, "DISPARO", "mira no cursor"),
        ("MERGULHADOR", "Mobilidade", CYAN, "DASH", "avanço invulneravel"),
    ]
    for i, (name, role, col, action, note) in enumerate(cards):
        x0 = 85 + i*510
        d.rounded_rectangle((x0, 55, x0+450, 665), radius=34, fill="#FFFFFF", outline=f"#{col}", width=8)
        d.rectangle((x0+155, 125, x0+295, 355), fill=f"#{col}", outline=f"#{NAVY}", width=6)
        d.ellipse((x0+180, 78, x0+270, 168), fill=f"#{col}", outline=f"#{NAVY}", width=6)
        d.text((x0+32, 395), name, font=fit_font(32, True), fill=f"#{NAVY}")
        d.text((x0+32, 445), role, font=fit_font(24), fill=f"#{MUTED}")
        d.rounded_rectangle((x0+30, 505, x0+420, 615), radius=20, fill=f"#{NAVY}")
        d.text((x0+55, 525), action, font=fit_font(27, True), fill=f"#{col}")
        d.text((x0+55, 568), note, font=fit_font(18), fill=f"#{WHITE}")
    img.save(GENERATED_DIR / "characters.png", quality=95)

    # Irregular isometric area diagram.
    img = Image.new("RGB", (1600, 850), "#07131F")
    d = ImageDraw.Draw(img)
    polygon = [(190, 390), (440, 190), (785, 235), (1020, 100), (1400, 260), (1320, 590), (1090, 730), (760, 650), (470, 760), (150, 610)]
    d.polygon(polygon, fill="#153D4B", outline=f"#{CYAN}")
    for cx, cy, col in [(460, 430, TEAL), (825, 420, PURPLE), (1130, 390, CORAL)]:
        draw_diamond(d, cx, cy, 430, 250, f"#{col}", f"#{GOLD}", 5)
    d.line((240, 560, 630, 460, 950, 500, 1320, 360), fill=f"#{GOLD}", width=18)
    d.ellipse((380, 370, 420, 410), fill=f"#{WHITE}", outline=f"#{NAVY}", width=3)
    for x, y in [(560, 330), (690, 520), (850, 315), (980, 580), (1180, 300), (1240, 520), (780, 240), (420, 600)]:
        d.rectangle((x-18, y-18, x+18, y+18), fill=f"#{CORAL}")
    d.text((205, 85), "AREA INICIAL - PROTOTIPO", font=fit_font(42, True), fill=f"#{WHITE}")
    d.text((210, 145), "tres setores conectados, contorno irregular e oito encontros", font=fit_font(24), fill=f"#{LIGHT}")
    img.save(GENERATED_DIR / "initial_area.png", quality=95)

    # Color palette.
    img = Image.new("RGB", (1600, 440), f"#{PAPER}")
    d = ImageDraw.Draw(img)
    swatches = [("ABISMO", NAVY), ("RUINA", DEEP), ("ATLANTIS", TEAL), ("BIOLUMIN.", CYAN), ("OURO", GOLD), ("CORAL", CORAL), ("MISTERIO", PURPLE)]
    x = 70
    for label, col in swatches:
        d.rounded_rectangle((x, 70, x+190, 290), radius=22, fill=f"#{col}", outline="#FFFFFF", width=4)
        d.text((x+12, 320), label, font=fit_font(19, True), fill=f"#{INK}")
        d.text((x+12, 354), f"#{col}", font=fit_font(17), fill=f"#{MUTED}")
        x += 215
    img.save(GENERATED_DIR / "palette.png", quality=95)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.58)
    sec.bottom_margin = Inches(0.58)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    sec.header_distance = Inches(0.25)
    sec.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (("Title", 32, NAVY), ("Heading 1", 22, NAVY), ("Heading 2", 15, TEAL), ("Heading 3", 11.5, PURPLE)):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    header = sec.header
    header.paragraphs[0].clear()
    footer = sec.footer
    p = footer.paragraphs[0]
    r = p.add_run("Documento vivo - prototipo Godot")
    set_run_font(r, size=8, color=MUTED)
    p.add_run(" " * 18)
    add_page_number(p)


def add_text(doc: Document, text: str, bold_prefix: str | None = None, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color or INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, italic=italic, color=color or INK)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color or INK)
    return p


def add_bullets(doc: Document, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.2, color=INK)


def add_numbered(doc: Document, items):
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{index}.  {item}")
        set_run_font(r, size=10.2, color=INK)


def add_section_title(doc: Document, number: str, title: str, subtitle: str = ""):
    if number in {"02", "04", "06", "08", "09", "10", "11", "12", "14", "15", "16"}:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        pad = 36 if number in {"04", "09", "11", "12", "16"} else 18
        spacer.paragraph_format.line_spacing = Pt(pad)
        r = spacer.add_run(" ")
        set_run_font(r, size=pad, color=WHITE)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_fixed_cell_width(table.cell(0, 0), 0.78)
    set_fixed_cell_width(table.cell(0, 1), 6.1)
    c0, c1 = table.rows[0].cells
    set_cell_shading(c0, GOLD)
    set_cell_shading(c1, NAVY)
    set_cell_margins(c0, 130, 100, 130, 100)
    set_cell_margins(c1, 130, 170, 130, 170)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(number)
    set_run_font(r, size=19, bold=True, color=NAVY)
    p = c1.paragraphs[0]
    r = p.add_run(title.upper())
    set_run_font(r, size=18, bold=True, color=WHITE)
    if subtitle:
        p = c1.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        r = p.add_run(subtitle)
        set_run_font(r, size=9, color=CYAN)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc: Document, title: str, body: str, accent=CYAN):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_fixed_cell_width(table.cell(0, 0), 0.14)
    set_fixed_cell_width(table.cell(0, 1), 6.75)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), PALE)
    set_cell_margins(table.cell(0, 1), 115, 150, 115, 150)
    p = table.cell(0, 1).paragraphs[0]
    r = p.add_run(title.upper())
    set_run_font(r, size=9.5, bold=True, color=accent)
    p = table.cell(0, 1).add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers, rows, widths=None, small=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        if widths:
            set_fixed_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        r = p.add_run(header)
        set_run_font(r, size=8.5 if small else 9, bold=True, color=WHITE)
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            if widths:
                set_fixed_cell_width(cells[i], widths[i])
            set_cell_shading(cells[i], "F2F7F6" if row_index % 2 == 0 else WHITE)
            set_cell_margins(cells[i], 80, 95, 80, 95)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_run_font(r, size=8.3 if small else 9.2, color=INK)
            edge = {"val": "single", "sz": "3", "color": "B6C9C8"}
            set_cell_border(cells[i], top=edge, bottom=edge, start=edge, end=edge)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc: Document, path: Path, width: float, caption: str, source: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(caption)
    set_run_font(r, size=8.3, italic=True, color=MUTED)
    if source:
        r = p.add_run(" Fonte: ")
        set_run_font(r, size=8.3, color=MUTED)
        add_hyperlink(p, source, source, color=TEAL)


def page_break(doc: Document):
    if not doc.paragraphs:
        paragraph = doc.add_paragraph()
    else:
        paragraph = doc.paragraphs[-1]
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_visual_assets()
    doc = Document()
    configure_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(GENERATED_DIR / "cover.png"), width=Inches(7.02))
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [("ESTADO", "Protótipo jogável"), ("VERSÃO", "0.1 - 19 ago 2026"), ("ENGINE", "Godot 4.x")]
    for i, (label, value) in enumerate(data):
        cell = table.cell(0, i)
        set_cell_shading(cell, NAVY if i != 1 else DEEP)
        set_fixed_cell_width(cell, 2.32)
        set_cell_margins(cell, 120, 120, 120, 120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        set_run_font(r, size=8, bold=True, color=CYAN)
        r = p.add_run(value)
        set_run_font(r, size=10, bold=True, color=WHITE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento de visão, gameplay, arte, narrativa, tecnologia e produção")
    set_run_font(r, size=9.5, color=MUTED)

    # TOC
    page_break(doc)
    add_section_title(doc, "00", "Índice do documento", "Estrutura baseada no modelo fornecido e expandida para o projeto atual")
    toc = [
        ("01", "Introdução e visão do jogo"), ("02", "Pilares e experiência do jogador"),
        ("03", "Ciclo de gameplay"), ("04", "Personagens jogáveis"),
        ("05", "Movimento, combate e controles"), ("06", "Mundo e level design"),
        ("07", "Narrativa e diálogos"), ("08", "Lore - espaço reservado"),
        ("09", "Arte e referências visuais"), ("10", "UI e experiência"),
        ("11", "Áudio"), ("12", "Arquitetura técnica"),
        ("13", "Estado atual do protótipo"), ("14", "Roadmap e MVP"),
        ("15", "Riscos e decisões em aberto"), ("16", "Fontes e créditos"),
    ]
    table = doc.add_table(rows=8, cols=2)
    table.autofit = False
    for idx, (num, title) in enumerate(toc):
        row, col = divmod(idx, 2)
        cell = table.cell(row, col)
        set_cell_shading(cell, "F2F7F6" if row % 2 == 0 else WHITE)
        set_cell_margins(cell, 120, 140, 120, 140)
        set_fixed_cell_width(cell, 3.45)
        p = cell.paragraphs[0]
        r = p.add_run(num + "  ")
        set_run_font(r, size=11, bold=True, color=GOLD)
        r = p.add_run(title)
        set_run_font(r, size=10, bold=True, color=NAVY)
    doc.add_paragraph()
    add_callout(doc, "Como usar este GDD", "As marcações IMPLEMENTADO e PLANEJADO separam o protótipo funcional das propostas de expansão. O documento deve ser atualizado quando uma decisão de design mudar ou quando uma etapa do roadmap for concluída.", TEAL)
    doc.add_heading("Controle de versão", level=2)
    add_table(doc, ["Versão", "Data", "Escopo", "Responsável"], [["0.1", "19/08/2026", "Consolidação do protótipo, referências e direção de produção", "Equipe Drowned"]], [0.8, 1.1, 3.7, 1.3])

    # 01
    page_break(doc)
    add_section_title(doc, "01", "Introdução e visão do jogo", "High concept, gênero, público, plataforma e proposta")
    doc.add_heading("Resumo do jogo", level=2)
    add_text(doc, "Drowned é um RPG de ação 2D em perspectiva isométrica ambientado em uma Atlântida submersa. Após o naufrágio, o jogador escolhe um entre três sobreviventes com estilos distintos, explora ruínas conectadas, enfrenta criaturas abissais e descobre relações e conflitos por meio de diálogos cinematográficos reutilizáveis.")
    add_callout(doc, "Elevator pitch", "Soul Knight Prequel encontra a mitologia submersa de Atlântida, com combate responsivo, três identidades jogáveis e apresentação cromática de alto contraste inspirada em Hades.", GOLD)
    doc.add_heading("Ficha rápida", level=2)
    add_table(doc, ["Item", "Definição atual"], [
        ["Título", "Drowned"], ["Gênero", "RPG de ação isométrico / aventura narrativa"],
        ["Modo", "Um jogador no escopo atual; cooperação é possibilidade futura, não compromisso"],
        ["Plataforma inicial", "PC (Windows), com arquitetura de input preparada para teclado, mouse e controle"],
        ["Engine", "Godot 4.x"], ["Câmera", "Top-down isométrica, com movimento em oito direções"],
        ["Estado", "Protótipo funcional com placeholders geométricos"],
    ], [1.7, 5.2])
    doc.add_heading("Público-alvo", level=2)
    add_bullets(doc, [
        "Jogadores que valorizam combate rápido, leitura clara de perigos e personagens com funções diferentes.",
        "Público interessado em mitologia, cidades perdidas, fantasia submarina e mistério.",
        "Faixa sugerida: adolescentes e adultos; classificação depende do tratamento final de violência e horror abissal.",
        "Sessões curtas de exploração e combate, conectadas por progressão narrativa mais longa.",
    ])
    doc.add_heading("Plataforma e ferramentas", level=2)
    add_text(doc, "A base é desenvolvida em Godot. Arte final poderá ser produzida em Aseprite, Krita ou ferramenta equivalente; áudio em DAW a definir. Os sistemas foram separados por componentes para permitir substituição dos retângulos por sprites sem reescrever a movimentação ou o combate.")

    # 02
    page_break(doc)
    add_section_title(doc, "02", "Pilares e experiência do jogador", "Princípios que orientam decisões de conteúdo e sistema")
    pillars = [
        ["1. Movimento legível", "Controle imediato em oito direções, aceleração mínima e câmera que favorece antecipação."],
        ["2. Identidades distintas", "Cada personagem resolve o mesmo espaço de forma diferente: aproximação, distância ou mobilidade."],
        ["3. Atlântida viva", "Ruínas, tecnologia antiga, fauna abissal e sinais de civilização contam história pelo cenário."],
        ["4. Combate expressivo", "Feedback visual, cooldowns claros, dano, recuo e janelas de risco tornam cada ação compreensível."],
        ["5. Narrativa encenada", "Diálogos com três slots, transições de elenco e barras cinematográficas conectam personagens e mundo."],
    ]
    add_table(doc, ["Pilar", "Aplicação"], pillars, [2.0, 4.9])
    doc.add_heading("Experiência pretendida", level=2)
    add_text(doc, "O jogador deve sentir curiosidade ao avançar por uma cidade impossível, domínio ao aprender seu personagem e tensão quando o cenário alterna beleza bioluminescente com ameaça abissal. A leitura do combate deve permanecer limpa mesmo quando a arte for detalhada.")
    doc.add_heading("Princípios de decisão", level=2)
    add_bullets(doc, [
        "Se uma decoração reduz a leitura de hitboxes, ela deve perder contraste ou sair do plano de gameplay.",
        "Uma habilidade nova precisa alterar posição, alvo ou ritmo; não apenas aumentar números.",
        "A narrativa deve entrar em momentos curtos e significativos, sem interromper repetidamente a exploração.",
        "O protótipo deve validar sensação e clareza antes de receber produção artística definitiva.",
    ])
    add_figure(doc, GENERATED_DIR / "core_loop.png", 6.75, "Figura 1 - Ciclo macro proposto para uma expedição.")

    # 03
    page_break(doc)
    add_section_title(doc, "03", "Ciclo de gameplay", "Do menu de seleção ao encerramento da área")
    doc.add_heading("Fluxo atual implementado", level=2)
    add_numbered(doc, [
        "Selecionar Quebra-Mar, Vigia ou Mergulhador.",
        "Assistir à introdução específica do personagem com os dois amigos.",
        "Ver os amigos saírem e a criatura abissal entrar no diálogo.",
        "Explorar a área inicial irregular com câmera e oito encontros.",
        "Usar a ação principal, administrar vida e eliminar inimigos.",
        "Receber vitória ao derrotar todos ou derrota ao perder toda a vida.",
        "Reiniciar a área ou retornar à seleção de personagem.",
    ])
    doc.add_heading("Ciclo planejado para a versão de jogo", level=2)
    add_table(doc, ["Momento", "Ação", "Recompensa / decisão"], [
        ["Preparação", "Escolher personagem, equipamento e objetivo", "Definir estratégia"],
        ["Entrada", "Diálogo ou evento contextual", "Motivação e informação"],
        ["Exploração", "Ler rotas, segredos, perigos e NPCs", "Recursos, lore e opções"],
        ["Confronto", "Combinar movimento, habilidade e posicionamento", "Progressão e domínio"],
        ["Consequência", "Escolha, resgate, perda ou descoberta", "Estado narrativo persistente"],
        ["Retorno", "Atualizar base e preparar nova incursão", "Build e relações"],
    ], [1.35, 3.0, 2.55])
    add_callout(doc, "Escopo protegido", "Loot, builds, base hub e escolhas persistentes são propostas de design. Antes de produção, cada sistema deve ser validado contra o tamanho da equipe e o cronograma.", CORAL)
    doc.add_heading("Condições atuais", level=2)
    add_table(doc, ["Condição", "Regra"], [["Vitória", "Todos os inimigos da área derrotados"], ["Derrota", "Vida do jogador chega a zero"], ["Reinício", "Recarrega a cena da área"], ["Troca", "Retorna à seleção e executa uma nova introdução"]], [1.5, 5.4])

    # 04
    page_break(doc)
    add_section_title(doc, "04", "Personagens jogáveis", "Três formas de atravessar a mesma Atlântida")
    add_figure(doc, GENERATED_DIR / "characters.png", 6.8, "Figura 2 - Placeholders e funções de combate já implementadas.")
    add_table(doc, ["Personagem", "Função", "Ação", "Leitura de jogo", "Estado"], [
        ["Quebra-Mar", "Vanguarda", "Golpe frontal de curto alcance", "Aproxima, controla espaço e aceita mais risco", "IMPLEMENTADO"],
        ["Vigia", "Atirador", "Projétil retangular direcionado ao cursor", "Mantém distância e escolhe linhas de tiro", "IMPLEMENTADO"],
        ["Mergulhador", "Mobilidade", "Dash com dano e invulnerabilidade temporária", "Reposiciona e atravessa janelas perigosas", "IMPLEMENTADO"],
    ], [1.15, 1.0, 2.0, 2.2, 0.8], small=True)
    doc.add_heading("Diretrizes para expansão", level=2)
    add_bullets(doc, [
        "Preservar silhuetas, cores de acento e efeitos próprios para reconhecer o personagem em movimento.",
        "Cada kit futuro deve conter ação básica, mobilidade, recurso defensivo e pelo menos uma decisão de build.",
        "Nomes e atributos atuais são provisórios; os IDs internos devem permanecer estáveis quando saves existirem.",
        "O jogador ocupa por padrão o slot central em diálogos; retratos finais devem ser de corpo inteiro e com transparência.",
    ])
    doc.add_heading("Animações provisórias", level=2)
    add_text(doc, "Os retângulos já piscam ou esticam durante golpe, disparo e dash. Esses mesmos gatilhos devem alimentar AnimatedSprite2D ou AnimationPlayer quando a arte final chegar, preservando duração, cooldown e janela de dano.")

    # 05
    page_break(doc)
    add_section_title(doc, "05", "Movimento, combate e controles", "Responsividade inspirada em action RPGs isométricos")
    doc.add_heading("Movimentação", level=2)
    add_text(doc, "O jogador se move em oito direções sobre um plano 2D, com velocidade normalizada para evitar ganho diagonal. A câmera acompanha o personagem e respeita os limites amplos do mundo. O contorno caminhável é um polígono irregular, não um quadrado de teste.")
    doc.add_heading("Controles atuais", level=2)
    add_table(doc, ["Entrada", "Função"], [
        ["WASD / setas", "Movimentação em oito direções"], ["Espaço / J / mouse esquerdo", "Ação principal do personagem"],
        ["Enter / Espaço / clique em diálogo", "Completar typewriter ou avançar fala"], ["Esc", "Retornar à seleção"],
        ["[DEBUG] Curar", "Restaurar vida por completo"], ["[DEBUG] Morrer", "Disparar derrota imediatamente"],
    ], [2.2, 4.7])
    doc.add_heading("Modelo de combate", level=2)
    add_table(doc, ["Sistema", "Comportamento"], [
        ["Vida", "Componente compartilhado com dano, cura, morte e sinais para HUD"],
        ["Hitbox / Hurtbox", "Camadas independentes entregam dano e recuo sem acoplar atores"],
        ["Cooldown", "Cada perfil define intervalo da ação e o HUD recebe o uso por sinal"],
        ["Recuo", "Jogador e inimigos recebem impulso como feedback e ferramenta espacial"],
        ["Inimigos", "Perseguem, param em alcance, atacam e reagem a dano"],
        ["Dash", "Invulnerabilidade somente durante a janela ativa"],
    ], [1.65, 5.25])
    doc.add_heading("Mecânicas propostas", level=2)
    add_bullets(doc, [
        "Pressão da água: zonas profundas alteram mobilidade, projéteis ou consumo de recurso.",
        "Correntes: vetores ambientais que empurram atores e criam rotas temporárias.",
        "Bioluminescência: luz revela passagens, inimigos ou símbolos por tempo limitado.",
        "Relíquias atlantes: modificadores de build ligados a tecnologia antiga e pactos narrativos.",
    ])

    # 06
    page_break(doc)
    add_section_title(doc, "06", "Mundo e level design", "Atlântida como espaço navegável, perigoso e narrativo")
    add_figure(doc, GENERATED_DIR / "initial_area.png", 6.8, "Figura 3 - Leitura abstrata da área inicial implementada.")
    doc.add_heading("Área inicial", level=2)
    add_text(doc, "O protótipo usa um mundo de aproximadamente 2300 x 1400 pixels, organizado em três setores conectados. A grade isométrica é desenhada apenas dentro do piso caminhável, enquanto caminhos e ruínas dão escala e direção. Oito inimigos distribuem o primeiro teste de combate pelo espaço.")
    doc.add_heading("Gramática de salas proposta", level=2)
    add_table(doc, ["Tipo", "Função", "Elemento atlante"], [
        ["Travessia", "Ritmo e orientação", "Pontes quebradas, correntes e colunas"],
        ["Combate", "Teste de build e espaço", "Praças inundadas e santuários"],
        ["Descoberta", "Lore e recursos", "Arquivos, mosaicos e observatórios"],
        ["Narrativa", "Encontro encenado", "Salões, estátuas e portais"],
        ["Chefe", "Clímax mecânico", "Templo, abismo ou máquina oceânica"],
        ["Refúgio", "Recuperação e escolha", "Bolsa de ar, jardim bioluminescente"],
    ], [1.2, 2.4, 3.3])

    # 07
    page_break(doc)
    add_section_title(doc, "07", "Narrativa e diálogos", "Sistema reutilizável para cenas, NPCs e eventos de gameplay")
    doc.add_heading("Abertura atual", level=2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)
    add_numbered(doc, [
        "O personagem escolhido aparece no slot central.",
        "Os outros dois personagens aparecem como amigos nos slots laterais.",
        "O trio conversa com texto revelado gradualmente.",
        "Os amigos deixam a cena com transições animadas.",
        "Uma criatura substitui um dos lados, restando jogador e monstro.",
        "A conversa termina, o mundo é retomado e a área jogável começa.",
    ])
    doc.add_heading("Contrato do sistema", level=2)
    add_table(doc, ["Recurso", "Capacidade"], [
        ["Elenco", "Quantidade ilimitada de atores registrados por sequência"],
        ["Palco", "Três slots visíveis: esquerda, centro e direita"],
        ["Transições", "enter, exit e replace, executadas antes de uma fala"],
        ["Retrato", "Texture2D, caminho res:// ou placeholder colorido"],
        ["Texto", "Typewriter configurável e avanço por teclado ou clique"],
        ["Cinema", "Barras pretas, escurecimento do fundo, foco no falante e caixa inferior"],
        ["Integração", "DialogueManager pausa o mundo, instancia o overlay e restaura o estado"],
    ], [1.5, 5.4])
    doc.add_heading("Extensões futuras", level=2)
    add_bullets(doc, [
        "Escolhas com condições e consequências persistentes.",
        "Localização, histórico, voz, auto-advance e skip controlado.",
        "Expressões e poses por ator, sem trocar o contrato básico do slot.",
        "Gatilhos por entrada de área, interação, inspeção, fim de combate e objetivos.",
    ])
    add_callout(doc, "Direção de apresentação", "A referência de Persona 5 é usada apenas como intenção de foco: enquadramento forte, barras cinematográficas, composição assimétrica e transições expressivas. A interface final deve desenvolver identidade própria de Drowned.", PURPLE)

    # 08 lore blank
    page_break(doc)
    add_section_title(doc, "08", "Lore - espaço reservado", "Seção destinada ao texto autoral futuro")
    add_callout(doc, "Não preencher nesta versão", "Esta área foi deixada intencionalmente aberta para o autor definir cânone, cronologia, personagens, facções e eventos. As perguntas abaixo são apenas guias e não estabelecem respostas.", GOLD)
    prompts = [
        "Qual foi a causa do naufrágio e por que os três sobreviventes chegaram juntos?",
        "Atlântida está abandonada, adormecida, em guerra ou ainda habitada?",
        "O que significa estar 'drowned' neste universo: morte, transformação, pacto ou pertencimento?",
        "Quem é a criatura da introdução e o que ela quer do jogador?",
        "Quais facções disputam tecnologia, memória, território ou salvação?",
        "Como a escolha de personagem altera relações, revelações e finais?",
    ]
    for i, prompt in enumerate(prompts, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i:02d}. {prompt}")
        set_run_font(r, size=10.5, bold=True, color=TEAL)
        for _ in range(3):
            p = doc.add_paragraph("________________________________________________________________________________")
            p.paragraph_format.space_after = Pt(2)
            p.runs[0].font.color.rgb = rgb(LIGHT)

    # 09 art
    page_break(doc)
    add_section_title(doc, "09", "Arte e referências visuais", "Pixel art isométrica, Atlântida e contraste dramático")
    doc.add_heading("Direção visual", level=2)
    add_text(doc, "A meta é uma arte 2D isométrica com personagens legíveis, volumes compactos e cenários ricos em pequenas histórias visuais. A Atlântida deve combinar arquitetura grega reinterpretada, tecnologia oceânica, erosão, coral e bioluminescência. A coloração usa fundos profundos e frios com acentos quentes e saturados para personagens, ataques e objetivos.")
    add_figure(doc, GENERATED_DIR / "palette.png", 6.8, "Figura 4 - Paleta inicial proposta para Drowned.")
    add_table(doc, ["Referência", "Extrair", "Evitar copiar"], [
        ["Soul Knight Prequel", "Escala isométrica, personagens compactos, leitura por biomas e densidade controlada", "Silhuetas, assets, interface e composição específicas"],
        ["Hades", "Contraste, acentos quentes, recortes fortes, hierarquia visual e drama", "Paleta idêntica, personagens, molduras e iconografia proprietária"],
        ["Atlântida / arte grega", "Motivos marítimos, colunas, vasos, mosaicos, criaturas e geometria", "Reconstituição literal ou estereótipo sem coerência funcional"],
    ], [1.45, 3.0, 2.45], small=True)
    add_figure(doc, ASSET_DIR / "reference_soul_knight_prequel.png", 4.35, "Figura 5 - Soul Knight Prequel: referência para mapa isométrico pixelado e leitura de biomas.", "https://prequel.chillyroom.com/et/article/base/6")

    page_break(doc)
    add_section_title(doc, "09B", "Moodboard das bases", "Referências de cor, mitologia e materialidade")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (path, title, note) in enumerate([
        (ASSET_DIR / "reference_hades.png", "Hades", "Contraste entre preto, vermelho, laranja e ouro; personagem recortado sobre ambiente dramático."),
        (ASSET_DIR / "reference_poseidon_skyphos.jpg", "Poseidon e o cavalo-marinho", "Motivo marítimo grego, linhas negras e terracota. Obra em domínio público do The Met."),
    ]):
        cell = table.cell(0, i)
        set_fixed_cell_width(cell, 3.4)
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(3.05), height=Inches(3.65) if i == 0 else Inches(2.55))
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_run_font(r, size=11, bold=True, color=NAVY)
        p = cell.add_paragraph(note)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_run_font(r, size=8.7, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fontes: ")
    set_run_font(r, size=8.3, color=MUTED)
    add_hyperlink(p, "Supergiant Games - Hades", "https://www.supergiantgames.com/games/hades/", color=TEAL)
    r = p.add_run("  |  ")
    set_run_font(r, size=8.3, color=MUTED)
    add_hyperlink(p, "The Met - Terracotta skyphos", "https://www.metmuseum.org/art/collection/search/250547", color=TEAL)
    doc.add_heading("Aplicação no jogo", level=2)
    add_bullets(doc, [
        "Chão e estruturas: azul abissal, verde-petróleo, pedra úmida e baixo contraste.",
        "Interativos: ouro atlante, ciano bioluminescente e bordas mais nítidas.",
        "Ameaças: coral, magenta e roxo, reservados para dano, corrupção e chefes.",
        "Personagens: silhueta clara, 1 cor dominante e 1 acento de habilidade.",
        "FX: ataques devem ultrapassar temporariamente a luminosidade do cenário, sem esconder telegráficos.",
    ])

    # 10 UI
    page_break(doc)
    add_section_title(doc, "10", "UI e experiência", "Informação rápida sem competir com a atmosfera")
    doc.add_heading("Telas e HUD atuais", level=2)
    add_table(doc, ["Interface", "Conteúdo", "Estado"], [
        ["Seleção", "Três cartões gerados por dados, foco e confirmação", "IMPLEMENTADO"],
        ["Diálogo", "Três slots, barras, caixa, typewriter e transições", "IMPLEMENTADO"],
        ["HUD", "Vida, ação/cooldown, inimigos restantes e debug", "IMPLEMENTADO"],
        ["Fim de rodada", "Vitória/derrota, reinício e troca de personagem", "IMPLEMENTADO"],
        ["Inventário / build", "Equipamentos, relíquias e comparações", "PLANEJADO"],
        ["Mapa / objetivos", "Orientação, regiões e descobertas", "PLANEJADO"],
    ], [1.55, 3.75, 1.6])
    doc.add_heading("Diretrizes de UI", level=2)
    add_bullets(doc, [
        "Usar molduras atlantes simplificadas, sem reduzir área útil ou legibilidade.",
        "Manter texto sobre planos sólidos ou muito escurecidos; nunca diretamente sobre ruído do cenário.",
        "Diferenciar informação persistente, alerta temporário e debug por posição e cor.",
        "Garantir navegação por teclado e controle desde a primeira versão da tela.",
        "Oferecer ajuste de velocidade do texto, escala da interface, tremor, flashes e contraste.",
    ])
    doc.add_heading("Acessibilidade planejada", level=2)
    add_table(doc, ["Necessidade", "Solução"], [
        ["Daltonismo", "Ícones e formas além de cor para dano, raridade e personagem"],
        ["Sensibilidade visual", "Reduzir flashes, tremor e partículas"],
        ["Leitura", "Tamanho de texto, velocidade e fundo da caixa ajustáveis"],
        ["Entrada", "Remapeamento completo e suporte a controle"],
        ["Dificuldade", "Assistências separadas para dano recebido, velocidade e telegráficos"],
    ], [2.0, 4.9])

    # 11 audio
    page_break(doc)
    add_section_title(doc, "11", "Áudio", "Identidade submarina, feedback e silêncio dramático")
    doc.add_heading("Música", level=2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)
    add_text(doc, "A trilha deve misturar percussão profunda, metais processados, vozes distantes, texturas aquáticas e instrumentos que remetam ao Mediterrâneo sem imitação literal. Exploração pede camadas lentas e espaçosas; combate adiciona pulso e ataques rítmicos; diálogos importantes reduzem densidade para favorecer voz e leitura.")
    doc.add_heading("Efeitos sonoros", level=2)
    add_table(doc, ["Família", "Intenção"], [
        ["Movimento", "Passos amortecidos, água, equipamento e diferença de material"],
        ["Golpe", "Impacto curto, deslocamento de água e confirmação de acerto"],
        ["Disparo", "Carga cristalina, trajetória e impacto distinguíveis"],
        ["Dash", "Sucção, passagem e estalo final de reposicionamento"],
        ["Inimigos", "Assinaturas por espécie, telegráficos e estado de dano"],
        ["Ambiente", "Correntes, ruínas, fauna, pressão, estruturas e silêncio"],
        ["UI", "Sons discretos com materialidade de pedra, vidro marinho e metal"],
    ], [1.45, 5.45])
    add_callout(doc, "Prioridade", "Antes de música final, produzir um kit mínimo de feedback: uso de habilidade, acerto, dano recebido, morte, cura, vitória, derrota, foco e confirmação de menu.", GOLD)

    # 12 tech
    page_break(doc)
    add_section_title(doc, "12", "Arquitetura técnica", "Base modular pronta para trocar placeholders por arte")
    doc.add_heading("Fluxo de cenas", level=2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)
    add_text(doc, "character_select.tscn -> intro_dialogue.tscn -> movement_lab.tscn. GameState persiste o perfil escolhido; DialogueManager executa sequências; a arena instancia jogador e inimigos e observa vitória/derrota.")
    doc.add_heading("Sistemas principais", level=2)
    add_table(doc, ["Sistema", "Responsabilidade", "Extensão prevista"], [
        ["GameState", "Perfis e seleção atual", "Save, flags e progressão"],
        ["PlayerController", "Movimento, ação e feedback", "Kits, animações e estados"],
        ["Health / Hitbox / Hurtbox", "Vida, dano, cura, recuo e colisão", "Resistências e efeitos"],
        ["IsometricArena", "Contorno, regiões, grade e limites", "TileMap, salas e navegação"],
        ["DialogueManager", "Pausa, overlay e sequência", "Escolhas, condições e localização"],
        ["MovementLab", "Spawn, HUD, objetivos e fluxo", "Diretor de encontros e transições"],
    ], [1.5, 2.7, 2.7], small=True)
    doc.add_heading("Princípios de implementação", level=2)
    add_bullets(doc, [
        "Dados ajustáveis ficam fora da composição de cena sempre que possível.",
        "Atores compartilham componentes de vida e colisão em vez de duplicar regras.",
        "Sinais comunicam HUD, morte, dano e ações para reduzir acoplamento.",
        "Arte substitui nós visuais, preservando contratos, nomes essenciais e colisões.",
        "Smoke tests validam os três kits, cura, morte, vitória e encenação do diálogo.",
    ])
    doc.add_heading("Camadas de colisão atuais", level=2)
    add_table(doc, ["Valor", "Conteúdo", "Detectado por"], [["1", "Corpo do jogador", "Corpos inimigos"], ["2", "Corpos inimigos", "Jogador e inimigos"], ["4", "Hurtbox do jogador", "Ataques inimigos"], ["8", "Hurtboxes inimigas", "Golpe, dash e projéteis"]], [0.8, 2.4, 3.7])

    # 13 status
    page_break(doc)
    add_section_title(doc, "13", "Estado atual do protótipo", "O que já funciona e o que ainda é placeholder")
    add_table(doc, ["Área", "Entregue", "Próxima maturidade"], [
        ["Estrutura", "Pastas escaláveis, cenas e scripts separados", "Dados como Resources e módulos de conteúdo"],
        ["Personagens", "3 perfis, seleção e kits distintos", "Arte, animação, balanceamento e progressão"],
        ["Movimento", "8 direções, câmera e limites irregulares", "Obstáculos, navegação e controle"],
        ["Combate", "Dano, cura, morte, recuo, cooldown e inimigos", "Telegráficos, variedades, loot e chefes"],
        ["Área", "3 setores, caminhos, ruínas e 8 spawns", "TileMap, salas modulares e conteúdo"],
        ["Narrativa", "3 slots, transições, typewriter e introdução", "Escolhas, condições, saves e localização"],
        ["UI", "Seleção, HUD, debug e fim de rodada", "Identidade final e acessibilidade"],
        ["QA", "Smoke tests de combate e diálogo", "Testes de regressão, performance e playtests"],
    ], [1.2, 2.9, 2.8], small=True)
    doc.add_heading("Validações conhecidas", level=2)
    add_bullets(doc, [
        "COMBAT_SMOKE_TEST_OK: perfis, habilidades, spawns, cura, morte e resultados.",
        "DIALOGUE_SMOKE_TEST_OK: três roteiros, slots, saída dos amigos, entrada do monstro e ciclo completo.",
        "Documentação técnica e guia de customização estão dentro de res://docs.",
    ])
    doc.add_heading("Limitações atuais", level=2)
    add_bullets(doc, [
        "Inimigos perseguem em linha reta e não contornam obstáculos internos.",
        "O mapa é desenhado por código; ainda não usa TileMap, sprites ou carregamento modular.",
        "Não há áudio, loot, inventário, progressão, saves ou escolhas narrativas persistentes.",
        "Todos os visuais de personagem, inimigo e efeitos são formas geométricas.",
    ])

    # 14 roadmap
    page_break(doc)
    add_section_title(doc, "14", "Roadmap e MVP", "Sequência recomendada de produção")
    add_table(doc, ["Fase", "Objetivo", "Critério de pronto", "Estado"], [
        ["P0 - Fundação", "Estrutura, seleção, movimento, combate e diálogo", "Fluxo completo com placeholders e testes", "CONCLUÍDA"],
        ["P1 - Vertical slice", "1 área artística, 3 personagens, 3 inimigos, 1 mini-chefe", "10-15 min representativos e estáveis", "PRÓXIMA"],
        ["P2 - Progressão", "Loot/relíquias, builds, save e retorno", "Duas expedições com escolhas de build", "PLANEJADA"],
        ["P3 - Narrativa", "Lore, facções, escolhas e consequências", "Arco curto completo para 1 personagem", "AGUARDA LORE"],
        ["P4 - Conteúdo", "Biomas, chefes, NPCs, áudio e variedade", "MVP de início ao fim", "PLANEJADA"],
        ["P5 - Polimento", "Acessibilidade, performance, balanceamento e QA", "Build candidata a lançamento", "PLANEJADA"],
    ], [1.15, 2.2, 2.7, 0.85], small=True)
    doc.add_heading("Próximo marco recomendado: vertical slice", level=2)
    add_numbered(doc, [
        "Definir uma pequena bíblia visual e produzir o tileset da área inicial.",
        "Substituir um personagem e um inimigo por arte e animações finais para validar pipeline.",
        "Adicionar obstáculos com navegação e telegráficos de ataque.",
        "Criar um mini-chefe que cobre movimento, distância e dash.",
        "Implementar o primeiro pacote de áudio e feedback de UI.",
        "Executar playtests curtos, medir clareza, dificuldade e preferência de personagem.",
    ])
    add_callout(doc, "Definição de MVP", "Um jogo curto, completo e estável com os três personagens, um arco narrativo fechado, progressão suficiente para builds distintas, múltiplas áreas e um confronto final. O número exato de biomas e chefes deve ser definido após o vertical slice.", CORAL)

    # 15 risks
    page_break(doc)
    add_section_title(doc, "15", "Riscos e decisões em aberto", "Questões que afetam escopo, identidade ou arquitetura")
    add_table(doc, ["Risco / decisão", "Impacto", "Mitigação / pergunta"], [
        ["Escopo de RPG", "Loot, builds e narrativa podem multiplicar conteúdo", "Validar um ciclo mínimo antes de ampliar sistemas"],
        ["Três protagonistas", "Triplica arte, animação e variações narrativas", "Compartilhar base técnica e priorizar diferenças de alto valor"],
        ["Isométrico 2D", "Colisão pode divergir da arte", "Guia de footprint, pivô e altura para todos os assets"],
        ["Referências fortes", "Risco de identidade derivativa", "Criar regras visuais próprias de Atlântida e motivos exclusivos"],
        ["Lore em aberto", "Bloqueia facções, objetivos e progressão", "Preencher seção 08 antes de produção narrativa extensa"],
        ["IA simples", "Ambientes complexos quebram perseguição", "Adicionar NavigationRegion2D e testes por sala"],
        ["Cooperação", "Altera câmera, rede, balanceamento e interface", "Decidir antes do vertical slice; padrão atual é single-player"],
    ], [1.7, 2.4, 2.75], small=True)
    doc.add_heading("Decisões necessárias após este GDD", level=2)
    add_bullets(doc, [
        "Estrutura final: campanha linear, runs, hub ou combinação.",
        "Modelo de progressão: equipamento, árvore, relíquias, meta-progressão ou apenas habilidades.",
        "Tom da lore: aventura, tragédia, horror abissal, fantasia heroica ou mistura definida.",
        "Escopo de plataforma e suporte a controle no primeiro lançamento.",
        "Quantidade-alvo de áreas, inimigos, chefes e duração do MVP.",
    ])

    # 16 sources
    page_break(doc)
    add_section_title(doc, "16", "Fontes e créditos", "Referências externas e documentação interna")
    doc.add_heading("Referências externas", level=2)
    sources = [
        ("Soul Knight Prequel - site oficial da ChillyRoom", "https://prequel.chillyroom.com/et/article/base/6", "Pixel art isométrica, mapa, biomas e descrição oficial de action RPG."),
        ("Hades - site oficial da Supergiant Games", "https://www.supergiantgames.com/games/hades/", "Arte promocional, contraste, mitologia e apresentação visual."),
        ("Hades FAQ - Supergiant Games", "https://www.supergiantgames.com/blog/hades-faq/", "Descrição oficial de roguelike, ação, atmosfera e narrativa."),
        ("Terracotta skyphos - The Metropolitan Museum of Art", "https://www.metmuseum.org/art/collection/search/250547", "Poseidon em cavalo-marinho; referência marítima grega em domínio público."),
    ]
    for title, url, purpose in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_hyperlink(p, title, url, color=TEAL)
        p = doc.add_paragraph(purpose)
        p.paragraph_format.left_indent = Inches(0.22)
        for r in p.runs:
            set_run_font(r, size=9.2, color=MUTED)
    doc.add_heading("Documentação do projeto", level=2)
    add_bullets(doc, [
        "README.md - índice geral do projeto.",
        "docs/PROJECT_STRUCTURE.md - organização e convenções.",
        "docs/PLAYABLE_PROTOTYPE.md - fluxo e controles atuais.",
        "docs/TECHNICAL_REFERENCE.md - arquitetura, sinais, colisões e funções.",
        "docs/CUSTOMIZATION_GUIDE.md - receitas de alteração do protótipo.",
    ])
    doc.add_heading("Nota de uso das referências", level=2)
    add_text(doc, "As imagens externas são usadas como referência visual e permanecem propriedade de seus respectivos autores, exceto a obra do The Met indicada como domínio público. O objetivo é orientar linguagem e contraste, não reproduzir assets, personagens ou interface. Os diagramas Drowned deste documento foram criados especificamente para o projeto.")
    add_callout(doc, "Documento vivo", "Atualize versão, estado de implementação e decisões em aberto a cada marco. Quando a lore for escrita, substitua a página reservada e revise todas as seções dependentes: objetivos, facções, biomas, inimigos, diálogos e final.", GOLD)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
